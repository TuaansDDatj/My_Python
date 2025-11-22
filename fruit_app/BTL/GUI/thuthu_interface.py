import tkinter as tk
from tkinter import ttk
import tkinter.messagebox as messagebox
from datetime import datetime, timedelta

class ThuthuApp:
    def __init__(self, root, db, book_manager, user_manager, borrow_manager, return_manager, report_manager):
        self.root = root
        self.db = db
        self.book_manager = book_manager
        self.user_manager = user_manager
        self.borrow_manager = borrow_manager
        self.return_manager = return_manager
        self.report_manager = report_manager

        self.root.title("ỨNG DỤNG QUẢN LÝ THƯ VIỆN ĐẠI HỌC - THỦ THƯ")
        self.root.geometry("1050x580")
        self.root.configure(bg="#f8f9fa")
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        style = ttk.Style()
        style.theme_use('clam')

        style.configure('TButton', font=('Arial', 10), borderwidth=1)
        style.configure('Primary.TButton', background='#1abc9c', foreground='white')
        style.configure('Secondary.TButton', background='#3498db', foreground='white')
        style.configure('Warning.TButton', background='#f39c12', foreground='white')
        style.configure('Danger.TButton', background='#e74c3c', foreground='white')
        style.configure('Info.TButton', background='#9b59b6', foreground='white')

        style.configure('Treeview', font=('Arial', 9), rowheight=25)
        style.configure('Treeview.Heading', font=('Arial', 10, 'bold'))

        # Tạo khung
        self.create_menu_frame()
        self.content_frame = tk.Frame(root, bg="#f8f9fa")
        self.content_frame.pack(side="right", expand=True, fill="both")

        self.show_book_management()

    def print_borrow_slip(self, ma_phieu_muon):
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            import os
            from tkinter import filedialog

            borrow_info = self.borrow_manager.get_borrow_by_id(ma_phieu_muon)
            if not borrow_info:
                messagebox.showerror("Lỗi", "Không tìm thấy thông tin phiếu mượn!")
                return False

            details = self.borrow_manager.get_borrow_details(ma_phieu_muon)
            if not details:
                messagebox.showerror("Lỗi", "Không tìm thấy chi tiết phiếu mượn!")
                return False

            doc = Document()

            for section in doc.sections:
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)

            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header.add_run("PHIẾU MƯỢN SÁCH\n")
            header_run.font.size = Pt(16)
            header_run.font.bold = True

            slip_no = doc.add_paragraph()
            slip_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
            slip_no_run = slip_no.add_run(f"Mã phiếu: {ma_phieu_muon}\n")
            slip_no_run.font.size = Pt(12)
            slip_no_run.font.bold = True

            doc.add_paragraph(f"Mã độc giả: {borrow_info[1]}")
            doc.add_paragraph(f"Tên độc giả: {borrow_info[5]}")
            doc.add_paragraph(f"Ngày mượn: {borrow_info[2]}")
            doc.add_paragraph(f"Ngày hẹn trả: {borrow_info[3]}")
            doc.add_paragraph(f"Trạng thái: {borrow_info[4]}")

            doc.add_paragraph("_" * 50)

            doc.add_paragraph("CHI TIẾT SÁCH MƯỢN:", style='Heading 2')
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Set header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "STT"
            header_cells[1].text = "Mã Sách"
            header_cells[2].text = "Tên Sách"
            header_cells[3].text = "Số Lượng"
            header_cells[4].text = "Ghi Chú"

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add book details
            for idx, detail in enumerate(details, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = str(detail[2])  # MaSach
                row_cells[2].text = str(detail[5])  # TenSach
                row_cells[3].text = str(detail[3])  # SoLuong
                row_cells[4].text = ""  # Ghi Chú

            # Add signature sections
            doc.add_paragraph("\n\n")
            signatures = doc.add_paragraph()
            signatures.alignment = WD_ALIGN_PARAGRAPH.CENTER
            signatures.add_run("Người mượn\n(Ký và ghi rõ họ tên)").font.bold = True

            # Add current date at the bottom right
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_paragraph.add_run(
                f"\nNgày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}")

            # Open file dialog to choose save location and filename
            default_filename = f"PhieuMuon_{ma_phieu_muon}.docx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Lưu phiếu mượn"
            )

            # If user cancels the dialog, return
            if not file_path:
                return False

            # Save the document to the selected path
            doc.save(file_path)
            print(f"Tệp đã được lưu tại: {file_path}")

            # Open the document with the default application
            os.startfile(file_path)
            return True

        except ImportError as e:
            messagebox.showerror("Lỗi",
                                 f"Thiếu thư viện: {str(e)}\nHãy cài đặt python-docx bằng lệnh: pip install python-docx")
            return False
        except PermissionError:
            messagebox.showerror("Lỗi", "Không có quyền truy cập để lưu tệp. Vui lòng chọn vị trí lưu khác.")
            return False
        except Exception as e:
            messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")
            return False

    def print_return_slip(self, ma_phieu_tra):
        try:
            # Import necessary libraries
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            import os
            from tkinter import filedialog

            # Get return slip information
            cursor = self.db.conn.cursor()
            cursor.execute("""
                   SELECT pt.MaPhieuTra, pt.MaPhieuMuon, pt.NgayTra, 
                          pm.MaDocGia, nd.HoVaTen, pm.NgayMuon, pm.NgayHenTra
                   FROM PhieuTra pt
                   JOIN PhieuMuon pm ON pt.MaPhieuMuon = pm.MaPhieuMuon
                   JOIN NguoiDung nd ON pm.MaDocGia = nd.MaDocGia
                   WHERE pt.MaPhieuTra = ?
               """, (ma_phieu_tra,))

            return_info = cursor.fetchone()
            if not return_info:
                messagebox.showerror("Lỗi", "Không tìm thấy thông tin phiếu trả!")
                return False

            details = self.return_manager.get_return_details(ma_phieu_tra)
            if not details:
                messagebox.showerror("Lỗi", "Không tìm thấy chi tiết phiếu trả!")
                return False

            # Create a new Document
            doc = Document()

            # Set margins
            for section in doc.sections:
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)

            # Add header
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header.add_run("PHIẾU TRẢ SÁCH\n")
            header_run.font.size = Pt(16)
            header_run.font.bold = True

            # Add slip number
            slip_no = doc.add_paragraph()
            slip_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
            slip_no_run = slip_no.add_run(f"Mã phiếu: {ma_phieu_tra}\n")
            slip_no_run.font.size = Pt(12)
            slip_no_run.font.bold = True

            # Add reader information
            doc.add_paragraph(f"Mã độc giả: {return_info[3]}")
            doc.add_paragraph(f"Tên độc giả: {return_info[4]}")
            doc.add_paragraph(f"Mã phiếu mượn: {return_info[1]}")
            doc.add_paragraph(f"Ngày mượn: {return_info[5]}")
            doc.add_paragraph(f"Ngày hẹn trả: {return_info[6]}")
            doc.add_paragraph(f"Ngày trả: {return_info[2]}")

            # Add a line
            doc.add_paragraph("_" * 50)

            # Add book details table
            doc.add_paragraph("CHI TIẾT SÁCH TRẢ:", style='Heading 2')
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Set header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "STT"
            header_cells[1].text = "Mã Sách"
            header_cells[2].text = "Tên Sách"
            header_cells[3].text = "Số Lượng"
            header_cells[4].text = "Tình Trạng Sách"

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add book details
            for idx, detail in enumerate(details, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = str(detail[1])  # MaSach
                row_cells[2].text = str(detail[2])  # TenSach
                row_cells[3].text = str(detail[3])  # SoLuong
                row_cells[4].text = str(detail[4])  # TinhTrangSach

            # Add signature sections
            doc.add_paragraph("\n\n")
            signatures = doc.add_paragraph()
            signatures.alignment = WD_ALIGN_PARAGRAPH.CENTER
            signatures.add_run("Người trả\n(Ký và ghi rõ họ tên)").font.bold = True

            # Add current date at the bottom right
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_paragraph.add_run(
                f"\nNgày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}")

            # Open file dialog to choose save location and filename
            default_filename = f"PhieuTra_{ma_phieu_tra}.docx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Lưu phiếu trả"
            )

            # If user cancels the dialog, return
            if not file_path:
                return False

            # Save the document to the selected path
            doc.save(file_path)
            print(f"Tệp đã được lưu tại: {file_path}")

            # Open the document with the default application
            os.startfile(file_path)
            return True

        except ImportError as e:
            messagebox.showerror("Lỗi",
                                 f"Thiếu thư viện: {str(e)}\nHãy cài đặt python-docx bằng lệnh: pip install python-docx")
            return False
        except PermissionError:
            messagebox.showerror("Lỗi", "Không có quyền truy cập để lưu tệp. Vui lòng chọn vị trí lưu khác.")
            return False
        except Exception as e:
            messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")
            return False

    def create_menu_frame(self):
        menu_frame = tk.Frame(self.root, bg="#2c3e50", width=220)
        menu_frame.pack(side="left", fill="y")

        menu_frame.pack_propagate(False)

        # Tiêu đề và lời chào
        app_title = tk.Label(menu_frame, text="Thư Viện Đại Học", font=("Arial", 14, "bold"),
                             bg="#2c3e50", fg="white", justify="center")
        app_title.pack(pady=(20, 5))

        admin_text_label = tk.Label(menu_frame, text="Xin chào, Thủ Thư", font=("Arial", 12),
                                    bg="#2c3e50", fg="white", justify="center")
        admin_text_label.pack(pady=(0, 20))

        separator = ttk.Separator(menu_frame, orient='horizontal')
        separator.pack(fill='x', padx=10, pady=10)

        # Menu buttons
        self.create_menu_button(menu_frame, "📚 Quản lý Sách", self.show_book_management)
        self.create_menu_button(menu_frame, "📖 Mượn/Trả Sách", self.show_borrow_return_books)
        self.create_menu_button(menu_frame, "👥 Quản lý Người Dùng", self.show_users_management)
        #self.create_menu_button(menu_frame, "📊 Báo Cáo", self.show_reports)

        # Tạo phân cách
        separator = ttk.Separator(menu_frame, orient='horizontal')
        separator.pack(fill='x', padx=10, pady=10)

        # Thiết kế logout
        logout_btn = tk.Button(menu_frame, text="🚪 Đăng Xuất", font=("Arial", 12),
                               bg="#e74c3c", fg="white", relief="flat",
                               padx=10, pady=7, width=20, anchor="w",
                               activebackground="#c0392b", command=self.logout)
        logout_btn.pack(side="bottom", pady=20)

    def create_menu_button(self, parent, text, command):
        btn = tk.Button(parent, text=text, font=("Arial", 12),
                        bg="#34495e", fg="white", relief="flat",
                        padx=10, pady=7, width=20, anchor="w",
                        activebackground="#1abc9c", command=command)
        btn.pack(pady=5)
        return btn

    def show_book_management(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        lbl_title = tk.Label(self.content_frame, text="Quản lý Sách", font=("Arial", 16, "bold"), bg="#f8f9fa")
        lbl_title.pack(pady=10)

        # Tab sách và danh mục
        tab_control = ttk.Notebook(self.content_frame)
        books_tab = ttk.Frame(tab_control)
        categories_tab = ttk.Frame(tab_control)

        tab_control.add(books_tab, text='Sách')
        tab_control.add(categories_tab, text='Thể Loại')
        tab_control.pack(expand=1, fill='both', padx=10, pady=5)

        # === BOOKS TAB ===
        form_frame = tk.Frame(books_tab, bg="#f8f9fa")
        form_frame.pack(pady=10, fill="x", padx=10)

        labels = ["Mã Sách", "Mã Thể Loại", "Tên Sách", "Tác Giả", "Thể Loại", "Số Lượng", "Nhà Xuất Bản", "Giá Trị",
                  "Tình Trạng"]
        entries = {}

        # Create a grid layout for the form
        for i, label_text in enumerate(labels):
            # Tạo khung cho các trường
            field_frame = tk.Frame(form_frame, bg="#f8f9fa")
            field_frame.grid(row=i // 3, column=i % 3, padx=10, pady=5, sticky="w")

            # Add label
            label = tk.Label(field_frame, text=label_text, font=("Arial", 10), bg="#f8f9fa", width=12, anchor="w")
            label.pack(side="left")

            # Thêm tiện ích nhập liệu phù hợp
            if label_text == "Mã Thể Loại":
                # Tạo combobox cho thể loại
                categories = self.book_manager.get_all_categories()
                category_codes = [cat[0] for cat in categories]
                entry = ttk.Combobox(field_frame, values=category_codes, width=15)
                entry.pack(side="left", fill="x", expand=True)
            elif label_text == "Thể Loại":
                # Combobox chỉ đọc cho tên thể loại
                categories = self.book_manager.get_all_categories()
                category_names = [cat[1] for cat in categories]
                entry = ttk.Combobox(field_frame, values=category_names, state="readonly", width=15)
                entry.pack(side="left", fill="x", expand=True)
            elif label_text == "Tình Trạng":
                # Tạo combobox cho tình trạng
                entry = ttk.Combobox(field_frame, values=["Mới", "Tốt", "Bình thường", "Cũ", "Hư hỏng"], width=15)
                entry.pack(side="left", fill="x", expand=True)
            else:
                entry = tk.Entry(field_frame, width=18)
                entry.pack(side="left", fill="x", expand=True)
            entries[label_text] = entry

        # Liên kết Mã Thể Loại và Tên Thể Loại
        def update_category_name(event):
            selected_code = entries["Mã Thể Loại"].get()
            categories = self.book_manager.get_all_categories()
            for cat in categories:
                if cat[0] == selected_code:
                    entries["Thể Loại"].set(cat[1])
                    break

        entries["Mã Thể Loại"].bind("<<ComboboxSelected>>", update_category_name)

        # Khung cho button
        button_frame = tk.Frame(books_tab, bg="#f8f9fa")
        button_frame.pack(pady=10)

        # Biến để lưu trữ sách đã chọn
        self.selected_book = None

        # Chức năng xóa mục nhập
        def clear_entries():
            for entry in entries.values():
                if hasattr(entry, 'delete'):
                    entry.delete(0, tk.END)
                elif hasattr(entry, 'set'):
                    entry.set('')
            self.selected_book = None

        # Chức năng reload danh sách
        def refresh_book_list():
            for item in book_list.get_children():
                book_list.delete(item)

            books = self.book_manager.get_all_books()

            for book in books:
                try:
                    book_list.insert("", "end", values=(
                        book[0],  # MaSach
                        book[1],  # MaTheLoai
                        book[2],  # TenSach
                        book[3],  # TacGia
                        book[4],  # TenTheLoai
                        book[5],  # SoLuong
                        book[6],  # NhaXuatBan
                        f"{book[7]:,.0f}",  # GiaTri (formatted)
                        book[8]  # TinhTrang
                    ))
                except Exception as e:
                    print(f"Lỗi khi xử lý sách {book}: {str(e)}")

        # Add book function
        def add_book():
            try:
                ma_sach = entries["Mã Sách"].get().strip()
                ma_the_loai = entries["Mã Thể Loại"].get().strip()
                ten_sach = entries["Tên Sách"].get().strip()
                tac_gia = entries["Tác Giả"].get().strip()
                so_luong = entries["Số Lượng"].get().strip()
                nha_xuat_ban = entries["Nhà Xuất Bản"].get().strip()
                gia_tri = entries["Giá Trị"].get().strip()
                tinh_trang = entries["Tình Trạng"].get().strip()

                # Validate data
                if not all([ma_sach, ma_the_loai, ten_sach, tac_gia, so_luong, nha_xuat_ban, gia_tri, tinh_trang]):
                    messagebox.showwarning("Cảnh báo", "Vui lòng điền đầy đủ thông tin!")
                    return

                # Validate quantity
                if not so_luong.isdigit() or int(so_luong) < 0:
                    messagebox.showwarning("Cảnh báo", "Số lượng phải là số nguyên dương!")
                    return

                # Validate value
                try:
                    gia_tri_float = float(gia_tri.replace(',', ''))
                    if gia_tri_float < 0:
                        raise ValueError
                except ValueError:
                    messagebox.showwarning("Cảnh báo", "Giá trị phải là số dương!")
                    return

                # Add book to database
                if self.book_manager.add_book(ma_sach, ma_the_loai, ten_sach, tac_gia, int(so_luong),
                                              nha_xuat_ban, gia_tri_float, tinh_trang):
                    messagebox.showinfo("Thành công", "Thêm sách thành công!")
                    clear_entries()
                    refresh_book_list()
                else:
                    messagebox.showerror("Lỗi", "Không thể thêm sách. Mã sách có thể đã tồn tại.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Edit book function
        def edit_book():
            if not self.selected_book:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn sách cần sửa!")
                return

            try:
                ma_sach = entries["Mã Sách"].get().strip()
                ma_the_loai = entries["Mã Thể Loại"].get().strip()
                ten_sach = entries["Tên Sách"].get().strip()
                tac_gia = entries["Tác Giả"].get().strip()
                so_luong = entries["Số Lượng"].get().strip()
                nha_xuat_ban = entries["Nhà Xuất Bản"].get().strip()
                gia_tri = entries["Giá Trị"].get().strip()
                tinh_trang = entries["Tình Trạng"].get().strip()

                # Validate data
                if not all([ma_sach, ma_the_loai, ten_sach, tac_gia, so_luong, nha_xuat_ban, gia_tri, tinh_trang]):
                    messagebox.showwarning("Cảnh báo", "Vui lòng điền đầy đủ thông tin!")
                    return

                # Validate quantity
                if not so_luong.isdigit() or int(so_luong) < 0:
                    messagebox.showwarning("Cảnh báo", "Số lượng phải là số nguyên dương!")
                    return

                # Validate value
                try:
                    gia_tri_float = float(gia_tri.replace(',', ''))
                    if gia_tri_float < 0:
                        raise ValueError
                except ValueError:
                    messagebox.showwarning("Cảnh báo", "Giá trị phải là số dương!")
                    return

                # Update book
                if self.book_manager.update_book(ma_sach, ma_the_loai, ten_sach, tac_gia, int(so_luong),
                                                 nha_xuat_ban, gia_tri_float, tinh_trang):
                    messagebox.showinfo("Thành công", "Cập nhật sách thành công!")
                    clear_entries()
                    refresh_book_list()
                else:
                    messagebox.showerror("Lỗi", "Không thể cập nhật thông tin sách.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Delete book function
        def delete_book():
            if not self.selected_book:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn sách cần xóa!")
                return

            if messagebox.askyesno("Xác nhận", f"Bạn có chắc chắn muốn xóa sách '{self.selected_book[2]}' không?"):
                if self.book_manager.delete_book(self.selected_book[0]):
                    messagebox.showinfo("Thành công", "Xóa sách thành công!")
                    clear_entries()
                    refresh_book_list()
                else:
                    messagebox.showerror("Lỗi", "Không thể xóa sách. Sách có thể đang được mượn.")

        # Search book function
        def search_book():
            search_term = entries["Tên Sách"].get().strip()
            if not search_term:
                refresh_book_list()
                return

            for item in book_list.get_children():
                book_list.delete(item)

            books = self.book_manager.search_books(search_term)
            if not books:
                print("Không tìm thấy sách nào.")
                return

            for book in books:
                book_list.insert("", "end", values=(
                    book[0],  # MaSach
                    book[1],  # MaTheLoai
                    book[2],  # TenSach
                    book[3],  # TacGia
                    book[4],  # TenTheLoai
                    book[5],  # SoLuong
                    book[6],  # NhaXuatBan
                    f"{book[7]:,.0f}",  # GiaTri (formatted)
                    book[8]  # TinhTrang
                ))

        # Add function buttons
        btn_add = tk.Button(button_frame, text="Thêm", bg="#1abc9c", fg="white", font=("Arial", 10), width=10,
                            command=add_book)
        btn_add.pack(side="left", padx=5)

        btn_edit = tk.Button(button_frame, text="Sửa", bg="#f39c12", fg="white", font=("Arial", 10), width=10,
                             command=edit_book)
        btn_edit.pack(side="left", padx=5)

        btn_delete = tk.Button(button_frame, text="Xóa", bg="#e74c3c", fg="white", font=("Arial", 10), width=10,
                               command=delete_book)
        btn_delete.pack(side="left", padx=5)

        btn_search = tk.Button(button_frame, text="Tìm kiếm", bg="#3498db", fg="white", font=("Arial", 10), width=10,
                               command=search_book)
        btn_search.pack(side="left", padx=5)

        btn_clear = tk.Button(button_frame, text="Làm mới", bg="#95a5a6", fg="white", font=("Arial", 10), width=10,
                              command=clear_entries)
        btn_clear.pack(side="left", padx=5)

        # Create book list table
        columns = labels
        book_list = ttk.Treeview(books_tab, columns=columns, show="headings", height=15)

        for col in columns:
            book_list.heading(col, text=col)
            if col in ["Mã Sách", "Mã Thể Loại"]:
                book_list.column(col, width=80, anchor="center")
            elif col in ["Số Lượng", "Giá Trị"]:
                book_list.column(col, width=80, anchor="e")
            elif col in ["Tên Sách", "Tác Giả", "Nhà Xuất Bản"]:
                book_list.column(col, width=150)
            else:
                book_list.column(col, width=100)

        # Add scrollbar
        table_frame = tk.Frame(books_tab)
        table_frame.pack(fill="both", expand=True, padx=10, pady=5)

        y_scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=book_list.yview)
        book_list.configure(yscrollcommand=y_scrollbar.set)

        book_list.pack(side="left", fill="both", expand=True)
        y_scrollbar.pack(side="right", fill="y")

        # Handle row selection
        def on_book_selected(event):
            selected_items = book_list.selection()
            if selected_items:
                selected_item = selected_items[0]
                values = book_list.item(selected_item, "values")

                self.selected_book = values
                # Fill form fields
                for i, label in enumerate(labels):
                    if i < len(values):
                        if hasattr(entries[label], 'delete'):
                            entries[label].delete(0, tk.END)
                            entries[label].insert(0, values[i])
                        elif hasattr(entries[label], 'set'):
                            entries[label].set(values[i])

        book_list.bind("<<TreeviewSelect>>", on_book_selected)

        # === CATEGORIES TAB ===
        cat_form_frame = tk.Frame(categories_tab, bg="#f8f9fa")
        cat_form_frame.pack(pady=10, fill="x", padx=10)

        # Create fields for category management
        cat_labels = ["Mã Thể Loại", "Tên Thể Loại"]
        cat_entries = {}

        for i, label_text in enumerate(cat_labels):
            field_frame = tk.Frame(cat_form_frame, bg="#f8f9fa")
            field_frame.pack(pady=5, fill="x")

            label = tk.Label(field_frame, text=label_text, font=("Arial", 10), bg="#f8f9fa", width=12, anchor="w")
            label.pack(side="left", padx=5)

            entry = tk.Entry(field_frame, width=30)
            entry.pack(side="left", fill="x", expand=True, padx=5)
            cat_entries[label_text] = entry

        # Button frame for categories
        cat_button_frame = tk.Frame(categories_tab, bg="#f8f9fa")
        cat_button_frame.pack(pady=10)

        # Selected category
        self.selected_category = None

        # Function to clear category entries
        def clear_cat_entries():
            for entry in cat_entries.values():
                entry.delete(0, tk.END)
            self.selected_category = None

        # Function to refresh category list
        def refresh_category_list():
            for item in cat_list.get_children():
                cat_list.delete(item)

            categories = self.book_manager.get_all_categories()
            for cat in categories:
                cat_list.insert("", "end", values=(cat[0], cat[1]))

            # Also update category dropdowns in book tab
            categories = self.book_manager.get_all_categories()
            category_codes = [cat[0] for cat in categories]
            category_names = [cat[1] for cat in categories]
            entries["Mã Thể Loại"]['values'] = category_codes
            entries["Thể Loại"]['values'] = category_names

        # Add category function
        def add_category():
            try:
                ma_the_loai = cat_entries["Mã Thể Loại"].get().strip()
                ten_the_loai = cat_entries["Tên Thể Loại"].get().strip()

                if not all([ma_the_loai, ten_the_loai]):
                    messagebox.showwarning("Cảnh báo", "Vui lòng điền đầy đủ thông tin!")
                    return

                if self.book_manager.add_category(ma_the_loai, ten_the_loai):
                    messagebox.showinfo("Thành công", "Thêm thể loại thành công!")
                    clear_cat_entries()
                    refresh_category_list()
                else:
                    messagebox.showerror("Lỗi", "Không thể thêm thể loại. Mã thể loại có thể đã tồn tại.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Category buttons
        btn_add_cat = tk.Button(cat_button_frame, text="Thêm", bg="#1abc9c", fg="white", font=("Arial", 10), width=10,
                                command=add_category)
        btn_add_cat.pack(side="left", padx=5)

        btn_clear_cat = tk.Button(cat_button_frame, text="Làm mới", bg="#95a5a6", fg="white", font=("Arial", 10),
                                  width=10, command=clear_cat_entries)
        btn_clear_cat.pack(side="left", padx=5)

        # Create category list table
        cat_columns = ["Mã Thể Loại", "Tên Thể Loại"]
        cat_list = ttk.Treeview(categories_tab, columns=cat_columns, show="headings", height=15)

        for col in cat_columns:
            cat_list.heading(col, text=col)
            if col == "Mã Thể Loại":
                cat_list.column(col, width=100, anchor="center")
            else:
                cat_list.column(col, width=300)

        # Add scrollbar for category list
        cat_table_frame = tk.Frame(categories_tab)
        cat_table_frame.pack(fill="both", expand=True, padx=10, pady=5)

        cat_y_scrollbar = ttk.Scrollbar(cat_table_frame, orient="vertical", command=cat_list.yview)
        cat_list.configure(yscrollcommand=cat_y_scrollbar.set)

        cat_list.pack(side="left", fill="both", expand=True)
        cat_y_scrollbar.pack(side="right", fill="y")

        # Handle category selection
        def on_category_selected(event):
            selected_items = cat_list.selection()
            if selected_items:
                selected_item = selected_items[0]
                values = cat_list.item(selected_item, "values")

                self.selected_category = values
                # Fill category form fields
                cat_entries["Mã Thể Loại"].delete(0, tk.END)
                cat_entries["Mã Thể Loại"].insert(0, values[0])
                cat_entries["Tên Thể Loại"].delete(0, tk.END)
                cat_entries["Tên Thể Loại"].insert(0, values[1])

        cat_list.bind("<<TreeviewSelect>>", on_category_selected)

        # Load initial data
        refresh_book_list()
        refresh_category_list()

    def show_borrow_return_books(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Create notebook with tabs (chỉ giữ tab Mượn và tab Trả)
        tab_control = ttk.Notebook(self.content_frame)
        borrow_tab = ttk.Frame(tab_control)
        return_tab = ttk.Frame(tab_control)

        tab_control.add(borrow_tab, text='Mượn Sách')
        tab_control.add(return_tab, text='Trả Sách')
        tab_control.pack(expand=1, fill='both', padx=10, pady=10)

        # === BORROW TAB ===
        borrow_fields = ["Mã Độc Giả", "Tên Độc Giả", "Mã Sách", "Tên Sách", "Số Lượng", "Ngày Mượn", "Ngày Hẹn Trả"]
        form_frame_borrow = tk.Frame(borrow_tab, bg="#f8f9fa")
        form_frame_borrow.pack(pady=10, fill="x", padx=10)

        entries_borrow = {}
        for i, field_text in enumerate(borrow_fields):
            # Create frame for each field
            field_frame = tk.Frame(form_frame_borrow, bg="#f8f9fa")
            field_frame.grid(row=i // 3, column=i % 3, padx=10, pady=5, sticky="w")

            # Add label
            label = tk.Label(field_frame, text=field_text, font=("Arial", 10), bg="#f8f9fa", width=12, anchor="w")
            label.pack(side="left")

            if field_text == "Mã Độc Giả":
                # Combobox for reader ID
                users = self.user_manager.get_all_users()
                user_ids = [user[0] for user in users]
                entry = ttk.Combobox(field_frame, values=user_ids, width=15)
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Tên Độc Giả":
                # Read-only field for reader name
                entry = tk.Entry(field_frame, width=20, state="readonly")
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Mã Sách":
                # Combobox for book ID
                books = self.book_manager.get_all_books()
                book_ids = [book[0] for book in books]
                entry = ttk.Combobox(field_frame, values=book_ids, width=15)
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Tên Sách":
                # Read-only field for book title
                entry = tk.Entry(field_frame, width=25, state="readonly")
                entry.pack(side="left", fill="x", expand=True)
            elif field_text in ["Ngày Mượn", "Ngày Hẹn Trả"]:
                # Date fields with default values
                entry = tk.Entry(field_frame, width=15)
                today = datetime.now().strftime("%Y-%m-%d")
                if field_text == "Ngày Mượn":
                    entry.insert(0, today)
                else:  # Return due date
                    return_date = (datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d")
                    entry.insert(0, return_date)
                entry.pack(side="left", fill="x", expand=True)
            else:
                entry = tk.Entry(field_frame, width=15)
                entry.pack(side="left", fill="x", expand=True)
            entries_borrow[field_text] = entry

        # Link Reader ID to Reader Name
        def update_reader_name(event):
            selected_id = entries_borrow["Mã Độc Giả"].get()
            user = self.user_manager.get_user_by_id(selected_id)
            if user:
                entries_borrow["Tên Độc Giả"].configure(state="normal")
                entries_borrow["Tên Độc Giả"].configure(state="normal")
                entries_borrow["Tên Độc Giả"].delete(0, tk.END)
                entries_borrow["Tên Độc Giả"].insert(0, user[1])
                entries_borrow["Tên Độc Giả"].configure(state="readonly")

        entries_borrow["Mã Độc Giả"].bind("<<ComboboxSelected>>", update_reader_name)

        # Link Book ID to Book Name
        def update_book_name(event):
            selected_code = entries_borrow["Mã Sách"].get()
            book = self.book_manager.get_book_by_id(selected_code)
            if book:
                entries_borrow["Tên Sách"].configure(state="normal")
                entries_borrow["Tên Sách"].delete(0, tk.END)
                entries_borrow["Tên Sách"].insert(0, book[1])
                entries_borrow["Tên Sách"].configure(state="readonly")

        entries_borrow["Mã Sách"].bind("<<ComboboxSelected>>", update_book_name)

        # Button frame for borrowing
        button_frame_borrow = tk.Frame(borrow_tab, bg="#f8f9fa")
        button_frame_borrow.pack(pady=10)

        # Clear borrow form fields
        def clear_borrow_entries():
            for field, entry in entries_borrow.items():
                if field not in ["Ngày Mượn", "Ngày Hẹn Trả"]:
                    if hasattr(entry, 'delete'):
                        entry.delete(0, tk.END)
                    elif hasattr(entry, 'set'):
                        entry.set('')
                else:
                    # Update dates
                    today = datetime.now().strftime("%Y-%m-%d")
                    if field == "Ngày Mượn":
                        entry.delete(0, tk.END)
                        entry.insert(0, today)
                    else:  # Return date
                        return_date = (datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d")
                        entry.delete(0, tk.END)
                        entry.insert(0, return_date)

            # Reset read-only fields
            entries_borrow["Tên Độc Giả"].configure(state="normal")
            entries_borrow["Tên Độc Giả"].delete(0, tk.END)
            entries_borrow["Tên Độc Giả"].configure(state="readonly")

            entries_borrow["Tên Sách"].configure(state="normal")
            entries_borrow["Tên Sách"].delete(0, tk.END)
            entries_borrow["Tên Sách"].configure(state="readonly")

        # Refresh borrow list
        def refresh_borrow_list():
            for item in borrow_list.get_children():
                borrow_list.delete(item)

            borrows = self.borrow_manager.get_borrow_slips()
            for borrow in borrows:
                borrow_list.insert("", "end", values=(
                    borrow[0],  # MaPhieuMuon
                    borrow[1],  # MaDocGia
                    borrow[5],  # HoVaTen (from JOIN)
                    borrow[2],  # NgayMuon
                    borrow[3],  # NgayHenTra
                    borrow[4]  # TrangThai
                ))

        # đã sửa
        # Add borrow function
        def add_borrow():
            try:
                ma_doc_gia = entries_borrow["Mã Độc Giả"].get().strip()
                ma_sach = entries_borrow["Mã Sách"].get().strip()
                so_luong_str = entries_borrow["Số Lượng"].get().strip()

                # Kiểm tra nếu so_luong_str là số nguyên dương
                if not so_luong_str or not so_luong_str.isdigit() or int(so_luong_str) <= 0:
                    messagebox.showwarning("Cảnh báo", "Số lượng phải là số nguyên dương!")
                    return

                so_luong = int(so_luong_str)
                ngay_muon = entries_borrow["Ngày Mượn"].get().strip()
                ngay_tra = entries_borrow["Ngày Hẹn Trả"].get().strip()

                # Xác thực dữ liệu
                if not all([ma_doc_gia, ma_sach, so_luong, ngay_muon, ngay_tra]):
                    messagebox.showwarning("Cảnh báo", "Vui lòng điền đầy đủ thông tin!")
                    return

                # Kiểm tra số lượng sách có đủ để mượn không
                available, message = self.borrow_manager.check_book_availability(ma_sach, so_luong)
                if not available:
                    messagebox.showwarning("Không thể mượn sách", message)
                    return

                # Tạo phiếu mượn
                ma_phieu_muon = self.borrow_manager.create_borrow_slip(ma_doc_gia, ngay_muon, ngay_tra)
                if ma_phieu_muon:
                    # Thêm chi tiết phiếu mượn
                    if self.borrow_manager.add_borrow_detail(ma_phieu_muon, ma_sach, so_luong):
                        messagebox.showinfo("Thành công", "Đã thêm phiếu mượn thành công!")
                        clear_borrow_entries()  # Xóa các trường nhập liệu
                        refresh_borrow_list()  # Làm mới danh sách phiếu mượn

                        # Cung cấp tùy chọn in phiếu
                        if messagebox.askyesno("In phiếu", "Bạn có muốn in phiếu mượn không?"):
                            try:
                                success = self.print_borrow_slip(ma_phieu_muon)
                                if success:
                                    messagebox.showinfo("Thành công", f"Đã in phiếu mượn #{ma_phieu_muon}")
                                else:
                                    messagebox.showerror("Lỗi", "Không thể in phiếu mượn.")
                            except Exception as e:
                                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi khi in: {str(e)}")
                    else:
                        # Quan trọng: Xóa phiếu mượn nếu thêm chi tiết thất bại
                        self.borrow_manager.delete_borrow_slip(ma_phieu_muon)
                        messagebox.showwarning("Cảnh báo", "Không thể thêm chi tiết phiếu mượn.")
                else:
                    messagebox.showerror("Lỗi", "Không thể tạo phiếu mượn.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Add borrow buttons
        btn_add_borrow = tk.Button(button_frame_borrow, text="Thêm", bg="#1abc9c", fg="white", font=("Arial", 10),
                                   width=10, command=add_borrow)
        btn_add_borrow.pack(side="left", padx=5)

        btn_clear_borrow = tk.Button(button_frame_borrow, text="Làm mới", bg="#95a5a6", fg="white", font=("Arial", 10),
                                     width=10, command=clear_borrow_entries)
        btn_clear_borrow.pack(side="left", padx=5)

        def print_selected_borrow():
            selected_items = borrow_list.selection()
            if not selected_items:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn phiếu mượn cần in!")
                return

            selected_item = selected_items[0]
            values = borrow_list.item(selected_item, "values")
            ma_phieu_muon = values[0]

            print("ID phiếu mượn đã chọn:", ma_phieu_muon)  # Dòng gỡ lỗi

            try:
                success = self.print_borrow_slip(ma_phieu_muon)
                if success:
                    messagebox.showinfo("Thành công", f"Đã in phiếu mượn #{ma_phieu_muon}")
                else:
                    messagebox.showerror("Lỗi", "Không thể in phiếu mượn.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi khi in: {str(e)}")

        # Update your print button's command
        btn_print = tk.Button(button_frame_borrow, text="In phiếu", bg="#9b59b6", fg="white",
                              font=("Arial", 10), width=12, command=print_selected_borrow)
        btn_print.pack(side="left", padx=5)

        # Create borrow list table
        columns_borrow = ["Mã Phiếu", "Mã Độc Giả", "Tên Độc Giả", "Ngày Mượn", "Ngày Hẹn Trả", "Trạng Thái"]
        borrow_list = ttk.Treeview(borrow_tab, columns=columns_borrow, show="headings", height=15)

        for col in columns_borrow:
            borrow_list.heading(col, text=col)
            if col in ["Mã Phiếu", "Mã Độc Giả"]:
                borrow_list.column(col, width=80, anchor="center")
            elif col in ["Ngày Mượn", "Ngày Hẹn Trả", "Trạng Thái"]:
                borrow_list.column(col, width=100, anchor="center")
            else:
                borrow_list.column(col, width=200)

        # Add scrollbar
        borrow_table_frame = tk.Frame(borrow_tab)
        borrow_table_frame.pack(fill="both", expand=True, padx=10, pady=10)

        borrow_y_scrollbar = ttk.Scrollbar(borrow_table_frame, orient="vertical", command=borrow_list.yview)
        borrow_list.configure(yscrollcommand=borrow_y_scrollbar.set)

        borrow_list.pack(side="left", fill="both", expand=True)
        borrow_y_scrollbar.pack(side="right", fill="y")

        def show_borrow_details(event):
            selected_items = borrow_list.selection()
            if selected_items:
                selected_item = selected_items[0]
                values = borrow_list.item(selected_item, "values")
                ma_phieu_muon = values[0]

                # Create new window for details
                details_window = tk.Toplevel(self.root)
                details_window.title(f"Chi tiết phiếu mượn #{ma_phieu_muon}")
                details_window.geometry("700x500")
                details_window.resizable(True, True)
                details_window.configure(bg="#f8f9fa")

                # Get borrow details
                details = self.borrow_manager.get_borrow_details(ma_phieu_muon)

                # Display information
                header_frame = tk.Frame(details_window, bg="#f8f9fa")
                header_frame.pack(pady=10, fill="x", padx=20)

                lbl_title = tk.Label(header_frame, text=f"Chi tiết phiếu mượn #{ma_phieu_muon}",
                                     font=("Arial", 14, "bold"), bg="#f8f9fa")
                lbl_title.pack(side="left", pady=10)

                reader_info = tk.Label(header_frame, text=f"Độc giả: {values[2]} ({values[1]})", font=("Arial", 12),
                                       bg="#f8f9fa")
                reader_info.pack(side="right", pady=10)

                date_frame = tk.Frame(details_window, bg="#f8f9fa")
                date_frame.pack(fill="x", padx=20)

                borrow_date = tk.Label(date_frame, text=f"Ngày mượn: {values[3]}", font=("Arial", 10), bg="#f8f9fa")
                borrow_date.pack(side="left", padx=5)

                return_date = tk.Label(date_frame, text=f"Ngày hẹn trả: {values[4]}", font=("Arial", 10), bg="#f8f9fa")
                return_date.pack(side="left", padx=20)

                status = tk.Label(date_frame, text=f"Trạng thái: {values[5]}", font=("Arial", 10), bg="#f8f9fa")
                status.pack(side="left", padx=5)

                # Create table for details
                columns = ["Mã Chi Tiết", "Mã Sách", "Tên Sách", "Số Lượng", "Đã Trả", "Còn Lại"]
                detail_list = ttk.Treeview(details_window, columns=columns, show="headings", height=10)

                for col in columns:
                    detail_list.heading(col, text=col)
                    if col in ["Mã Chi Tiết", "Mã Sách"]:
                        detail_list.column(col, width=80, anchor="center")
                    elif col in ["Số Lượng", "Đã Trả", "Còn Lại"]:
                        detail_list.column(col, width=80, anchor="center")
                    else:
                        detail_list.column(col, width=250)

                # Add data to table
                for detail in details:
                    remaining = detail[3] - detail[4]  # SoLuong - DaTra
                    detail_list.insert("", "end", values=(
                        detail[0],  # MaChiTiet
                        detail[2],  # MaSach
                        detail[5],  # TenSach (from JOIN)
                        detail[3],  # SoLuong
                        detail[4],  # DaTra
                        remaining  # ConLai
                    ))

                # Add scrollbar
                detail_frame = tk.Frame(details_window)
                detail_frame.pack(fill="both", expand=True, padx=20, pady=10)

                detail_y_scrollbar = ttk.Scrollbar(detail_frame, orient="vertical", command=detail_list.yview)
                detail_list.configure(yscrollcommand=detail_y_scrollbar.set)

                detail_list.pack(side="left", fill="both", expand=True)
                detail_y_scrollbar.pack(side="right", fill="y")

                # Close button
                btn_close = tk.Button(details_window, text="Đóng", command=details_window.destroy,
                                      bg="#3498db", fg="white", font=("Arial", 10), width=10)
                btn_close.pack(pady=15)

        borrow_list.bind("<Double-1>", show_borrow_details)

        # === RETURN TAB ===
        # Tạo frame chính để chứa tất cả nội dung của tab trả sách
        return_main_frame = tk.Frame(return_tab, bg="#f8f9fa")
        return_main_frame.pack(fill="both", expand=True)

        # Form nhập liệu trả sách
        return_fields = ["Mã Độc Giả", "Tên Độc Giả", "Mã Sách", "Tên Sách", "Số Lượng", "Ngày Trả", "Tình Trạng Sách"]
        form_frame_return = tk.Frame(return_main_frame, bg="#f8f9fa")
        form_frame_return.pack(pady=10, fill="x", padx=10)

        entries_return = {}
        for i, field_text in enumerate(return_fields):
            # Create frame for each field
            field_frame = tk.Frame(form_frame_return, bg="#f8f9fa")
            field_frame.grid(row=i // 3, column=i % 3, padx=10, pady=5, sticky="w")

            # Add label
            label = tk.Label(field_frame, text=field_text, font=("Arial", 10), bg="#f8f9fa", width=12, anchor="w")
            label.pack(side="left")

            if field_text == "Mã Độc Giả":
                # Combobox for reader ID
                users = self.user_manager.get_all_users()
                user_ids = [user[0] for user in users]
                entry = ttk.Combobox(field_frame, values=user_ids, width=15)
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Tên Độc Giả":
                # Read-only field for reader name
                entry = tk.Entry(field_frame, width=20, state="readonly")
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Mã Sách":
                # Read-only field for book ID (will be filled from active borrows list)
                entry = tk.Entry(field_frame, width=15, state="readonly")
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Tên Sách":
                # Read-only field for book title
                entry = tk.Entry(field_frame, width=25, state="readonly")
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Ngày Trả":
                # Date field with today's date
                entry = tk.Entry(field_frame, width=15)
                entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Tình Trạng Sách":
                # Combobox for book condition
                entry = ttk.Combobox(field_frame,
                                     values=["Tốt", "Bình thường", "Hư hỏng nhẹ", "Hư hỏng nặng", "Mất sách"], width=15)
                entry.set("Tốt")  # Default value
                entry.pack(side="left", fill="x", expand=True)
            else:
                entry = tk.Entry(field_frame, width=15)
                entry.pack(side="left", fill="x", expand=True)
            entries_return[field_text] = entry

        # Link Reader ID to Reader Name
        def update_return_reader_name(event):
            selected_id = entries_return["Mã Độc Giả"].get()
            user = self.user_manager.get_user_by_id(selected_id)
            if user:
                entries_return["Tên Độc Giả"].configure(state="normal")
                entries_return["Tên Độc Giả"].delete(0, tk.END)
                entries_return["Tên Độc Giả"].insert(0, user[1])
                entries_return["Tên Độc Giả"].configure(state="readonly")

                # Refresh active borrows for this reader
                refresh_active_borrows()

        entries_return["Mã Độc Giả"].bind("<<ComboboxSelected>>", update_return_reader_name)

        # Button frame for returning
        button_frame_return = tk.Frame(return_main_frame, bg="#f8f9fa")
        button_frame_return.pack(pady=10)

        # Variable to store selected borrow
        self.selected_borrow = None

        # Refresh active borrows list
        def refresh_active_borrows():
            ma_doc_gia = entries_return["Mã Độc Giả"].get().strip()

            for item in return_list.get_children():
                return_list.delete(item)

            active_borrows = self.return_manager.get_active_borrows(ma_doc_gia if ma_doc_gia else None)
            for borrow in active_borrows:
                return_list.insert("", "end", values=(
                    borrow[0],  # MaPhieuMuon
                    borrow[1],  # MaDocGia
                    borrow[2],  # HoVaTen
                    borrow[5],  # MaSach
                    borrow[6],  # TenSach
                    borrow[7],  # SoLuong
                    borrow[8],  # DaTra
                    borrow[9],  # ConLai
                    borrow[3],  # NgayMuon
                    borrow[4]  # NgayHenTra
                ))

        # Hàm load lịch sử trả sách
        def load_return_history():
            for item in return_history_list.get_children():
                return_history_list.delete(item)

            returns = self.return_manager.get_return_history()
            for ret in returns:
                days_overdue = ret[7] if ret[7] and ret[7] > 0 else 0
                return_history_list.insert("", "end", values=(
                    ret[0],  # MaPhieuTra
                    ret[1],  # MaPhieuMuon
                    ret[3],  # MaDocGia
                    ret[4],  # HoVaTen
                    ret[2],  # NgayTra
                    ret[5],  # NgayMuon
                    ret[6],  # NgayHenTra
                    f"{days_overdue:.0f} ngày" if days_overdue > 0 else "0 ngày"  # DaysOverdue
                ))

        # Process return function
        def process_return():
            try:
                ma_doc_gia = entries_return["Mã Độc Giả"].get().strip()
                ma_sach = entries_return["Mã Sách"].get().strip()
                so_luong = entries_return["Số Lượng"].get().strip()
                ngay_tra = entries_return["Ngày Trả"].get().strip()
                tinh_trang = entries_return["Tình Trạng Sách"].get().strip()

                # Validate data
                if not all([ma_doc_gia, ma_sach, so_luong, ngay_tra, tinh_trang]):
                    messagebox.showwarning("Cảnh báo", "Vui lòng điền đầy đủ thông tin!")
                    return

                # Validate quantity
                if not so_luong.isdigit() or int(so_luong) <= 0:
                    messagebox.showwarning("Cảnh báo", "Số lượng phải là số nguyên dương!")
                    return

                # Check if borrow has been selected
                if not self.selected_borrow:
                    messagebox.showwarning("Cảnh báo", "Vui lòng chọn phiếu mượn cần trả!")
                    return

                # Check if return quantity is not greater than remaining
                remaining = int(self.selected_borrow[7])  # ConLai
                if int(so_luong) > remaining:
                    messagebox.showwarning("Cảnh báo",
                                           f"Số lượng trả không được vượt quá số lượng còn lại ({remaining})!")
                    return

                # Create return slip
                ma_phieu_muon = self.selected_borrow[0]
                ma_phieu_tra = self.return_manager.create_return_slip(ma_phieu_muon, ngay_tra)

                if ma_phieu_tra:
                    if self.return_manager.add_return_detail(ma_phieu_tra, ma_sach, int(so_luong), tinh_trang):
                        messagebox.showinfo("Thành công", "Trả sách thành công!")
                        # Clear fields
                        entries_return["Mã Sách"].configure(state="normal")
                        entries_return["Mã Sách"].delete(0, tk.END)
                        entries_return["Mã Sách"].configure(state="readonly")

                        entries_return["Tên Sách"].configure(state="normal")
                        entries_return["Tên Sách"].delete(0, tk.END)
                        entries_return["Tên Sách"].configure(state="readonly")

                        entries_return["Số Lượng"].delete(0, tk.END)

                        # Refresh lists
                        refresh_active_borrows()
                        load_return_history()

                        # Print option
                        if messagebox.askyesno("In phiếu", "Bạn có muốn in phiếu trả không?"):
                            try:
                                success = self.print_return_slip(ma_phieu_tra)
                                if success:
                                    messagebox.showinfo("Thành công", f"Đã in phiếu trả #{ma_phieu_tra}")
                                else:
                                    messagebox.showerror("Lỗi", "Không thể in phiếu trả.")
                            except Exception as e:
                                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi khi in: {str(e)}")
                    else:
                        messagebox.showerror("Lỗi", "Không thể thêm chi tiết phiếu trả.")
                else:
                    messagebox.showerror("Lỗi", "Không thể tạo phiếu trả.")

            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Return buttons
        btn_return = tk.Button(button_frame_return, text="Trả sách", bg="#1abc9c", fg="white", font=("Arial", 10),
                               width=10, command=process_return)
        btn_return.pack(side="left", padx=5)

        btn_refresh = tk.Button(button_frame_return, text="Làm mới", bg="#95a5a6", fg="white", font=("Arial", 10),
                                width=10, command=refresh_active_borrows)
        btn_refresh.pack(side="left", padx=5)


        # Phần bảng danh sách phiếu mượn đang hoạt động
        # Tạo label cho bảng
        active_label = tk.Label(return_main_frame, text="Danh sách phiếu mượn đang hoạt động",
                                font=("Arial", 12, "bold"), bg="#f8f9fa")
        active_label.pack(pady=(5, 0), anchor="w", padx=10)

        # Frame cho bảng danh sách mượn đang hoạt động
        return_table_frame = tk.Frame(return_main_frame)
        return_table_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Return active borrows table
        columns_return = ["Mã Phiếu", "Mã Độc Giả", "Tên Độc Giả", "Mã Sách", "Tên Sách", "Số Lượng", "Đã Trả",
                          "Còn Lại", "Ngày Mượn", "Ngày Hẹn Trả"]
        return_list = ttk.Treeview(return_table_frame, columns=columns_return, show="headings", height=5)

        for col in columns_return:
            return_list.heading(col, text=col)
            if col in ["Mã Phiếu", "Mã Độc Giả", "Mã Sách", "Số Lượng", "Đã Trả", "Còn Lại"]:
                return_list.column(col, width=70, anchor="center")
            elif col in ["Ngày Mượn", "Ngày Hẹn Trả"]:
                return_list.column(col, width=100, anchor="center")
            else:
                return_list.column(col, width=150)

        # Add scrollbars
        return_y_scrollbar = ttk.Scrollbar(return_table_frame, orient="vertical", command=return_list.yview)
        return_list.configure(yscrollcommand=return_y_scrollbar.set)

        return_x_scrollbar = ttk.Scrollbar(return_table_frame, orient="horizontal", command=return_list.xview)
        return_list.configure(xscrollcommand=return_x_scrollbar.set)

        return_list.pack(side="top", fill="both", expand=True)
        return_y_scrollbar.pack(side="right", fill="y")
        return_x_scrollbar.pack(side="bottom", fill="x")

        # Phần lịch sử trả sách
        history_label = tk.Label(return_main_frame, text="Lịch sử trả sách",
                                 font=("Arial", 12, "bold"), bg="#f8f9fa")
        history_label.pack(pady=(15, 0), anchor="w", padx=10)

        # Thêm frame cho nút chức năng lịch sử
        history_button_frame = tk.Frame(return_main_frame, bg="#f8f9fa")
        history_button_frame.pack(pady=5, anchor="w", padx=10)

        # Hàm in phiếu trả từ lịch sử
        def print_history_return_slip():
            selected_items = return_history_list.selection()
            if not selected_items:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn một phiếu trả từ lịch sử!")
                return

            selected_item = selected_items[0]
            values = return_history_list.item(selected_item, "values")
            ma_phieu_tra = values[0]  # Mã phiếu trả

            try:
                success = self.print_return_slip(ma_phieu_tra)
                if success:
                    messagebox.showinfo("Thành công", f"Đã in phiếu trả #{ma_phieu_tra}")
                else:
                    messagebox.showerror("Lỗi", "Không thể in phiếu trả.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi khi in: {str(e)}")

        # Nút in phiếu trả từ lịch sử
        btn_print_history = tk.Button(history_button_frame, text="In phiếu đã chọn", bg="#9b59b6", fg="white",
                                      font=("Arial", 10), width=15, command=print_history_return_slip)
        btn_print_history.pack(side="left", padx=5)

        # Frame cho bảng lịch sử trả sách
        return_history_frame = tk.Frame(return_main_frame)
        return_history_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Return history table
        return_history_columns = ["Mã Phiếu Trả", "Mã Phiếu Mượn", "Mã Độc Giả", "Tên Độc Giả", "Ngày Trả", "Ngày Mượn",
                                  "Ngày Hẹn Trả", "Quá Hạn"]
        return_history_list = ttk.Treeview(return_history_frame, columns=return_history_columns, show="headings",
                                           height=10)

        for col in return_history_columns:
            return_history_list.heading(col, text=col)
            if col in ["Mã Phiếu Trả", "Mã Phiếu Mượn", "Mã Độc Giả"]:
                return_history_list.column(col, width=80, anchor="center")
            elif col in ["Ngày Trả", "Ngày Mượn", "Ngày Hẹn Trả", "Quá Hạn"]:
                return_history_list.column(col, width=100, anchor="center")
            else:
                return_history_list.column(col, width=150)

        # Add scrollbars
        history_y_scrollbar = ttk.Scrollbar(return_history_frame, orient="vertical", command=return_history_list.yview)
        return_history_list.configure(yscrollcommand=history_y_scrollbar.set)

        history_x_scrollbar = ttk.Scrollbar(return_history_frame, orient="horizontal",
                                            command=return_history_list.xview)
        return_history_list.configure(xscrollcommand=history_x_scrollbar.set)

        return_history_list.pack(side="top", fill="both", expand=True)
        history_y_scrollbar.pack(side="right", fill="y")
        history_x_scrollbar.pack(side="bottom", fill="x")

        def show_return_details(event):
            selected_items = return_list.selection()
            if selected_items:
                selected_item = selected_items[0]
                values = return_list.item(selected_item, "values")
                ma_phieu_tra = values[0]  # Giả sử đây là ID phiếu trả

                # Tạo cửa sổ mới để hiển thị chi tiết
                details_window = tk.Toplevel(self.root)
                details_window.title(f"Chi tiết phiếu trả #{ma_phieu_tra}")
                details_window.geometry("700x500")
                details_window.configure(bg="#f8f9fa")

                # Lấy chi tiết phiếu trả
                details = self.return_manager.get_return_details(ma_phieu_tra)
                print(details)
                # Hiển thị thông tin
                header_frame = tk.Frame(details_window, bg="#f8f9fa")
                header_frame.pack(pady=10, fill="x", padx=20)

                lbl_title = tk.Label(header_frame, text=f"Chi tiết phiếu trả #{ma_phieu_tra}",
                                     font=("Arial", 14, "bold"), bg="#f8f9fa")
                lbl_title.pack(side="left", pady=10)

                # Tạo bảng để hiển thị chi tiết
                columns = ["Mã Chi Tiết", "Mã Sách", "Tên Sách", "Số Lượng", "Tình Trạng"]
                detail_list = ttk.Treeview(details_window, columns=columns, show="headings", height=10)

                for col in columns:
                    detail_list.heading(col, text=col)
                    detail_list.column(col, width=80, anchor="center")

                # Thêm dữ liệu vào bảng
                for detail in details:
                    detail_list.insert("", "end", values=(
                        detail[0],  # Mã Chi Tiết
                        detail[1],  # Mã Sách
                        detail[2],  # Tên Sách
                        detail[3],  # Số Lượng
                        detail[4]  # Tình Trạng
                    ))

                # Thêm thanh cuộn
                detail_frame = tk.Frame(details_window)
                detail_frame.pack(fill="both", expand=True, padx=20, pady=10)

                # Tạo thanh cuộn dọc
                detail_y_scrollbar = ttk.Scrollbar(detail_frame, orient="vertical", command=detail_list.yview)
                detail_list.configure(yscrollcommand=detail_y_scrollbar.set)

                # Thêm bảng vào khung
                detail_list.pack(side="left", fill="both", expand=True)
                detail_y_scrollbar.pack(side="right", fill="y")
                # Nút Đóng
                btn_close = tk.Button(details_window, text="Đóng", command=details_window.destroy,
                                      bg="#3498db", fg="white", font=("Arial", 10), width=10)
                btn_close.pack(pady=15)

        return_list.bind("<Double-1>", show_return_details)

        # Hiển thị chi tiết khi double-click vào phiếu trả trong lịch sử
        def show_return_history_details(event):
            selected_items = return_history_list.selection()
            if selected_items:
                selected_item = selected_items[0]
                values = return_history_list.item(selected_item, "values")
                ma_phieu_tra = values[0]  # Mã phiếu trả

                # Hỏi người dùng có muốn in phiếu không
                if messagebox.askyesno("In phiếu trả", f"Bạn có muốn in phiếu trả #{ma_phieu_tra} không?"):
                    try:
                        success = self.print_return_slip(ma_phieu_tra)
                        if success:
                            messagebox.showinfo("Thành công", f"Đã in phiếu trả #{ma_phieu_tra}")
                        else:
                            messagebox.showerror("Lỗi", "Không thể in phiếu trả.")
                    except Exception as e:
                        messagebox.showerror("Lỗi", f"Đã xảy ra lỗi khi in: {str(e)}")
                    return

                # Create new window for details
                details_window = tk.Toplevel(self.root)
                details_window.title(f"Chi tiết phiếu trả #{ma_phieu_tra}")
                details_window.geometry("700x500")
                details_window.resizable(True, True)
                details_window.configure(bg="#f8f9fa")

                # Get return details
                details = self.return_manager.get_return_details(ma_phieu_tra)

                # Display information
                header_frame = tk.Frame(details_window, bg="#f8f9fa")
                header_frame.pack(pady=10, fill="x", padx=20)

                lbl_title = tk.Label(header_frame, text=f"Chi tiết phiếu trả #{ma_phieu_tra}",
                                     font=("Arial", 14, "bold"), bg="#f8f9fa")
                lbl_title.pack(side="left", pady=10)

                reader_info = tk.Label(header_frame, text=f"Phiếu mượn: #{values[1]}", font=("Arial", 12), bg="#f8f9fa")
                reader_info.pack(side="right", pady=10)

                date_frame = tk.Frame(details_window, bg="#f8f9fa")
                date_frame.pack(fill="x", padx=20)

                borrow_date = tk.Label(date_frame, text=f"Độc giả: {values[3]} ({values[2]})", font=("Arial", 10),
                                       bg="#f8f9fa")
                borrow_date.pack(side="left", padx=5)

                return_date = tk.Label(date_frame, text=f"Ngày trả: {values[4]}", font=("Arial", 10), bg="#f8f9fa")
                return_date.pack(side="left", padx=20)

                # Create table for details
                columns = ["Mã Chi Tiết", "Mã Sách", "Tên Sách", "Số Lượng", "Tình Trạng"]
                detail_list = ttk.Treeview(details_window, columns=columns, show="headings", height=10)

                for col in columns:
                    detail_list.heading(col, text=col)
                    if col in ["Mã Chi Tiết", "Mã Sách", "Số Lượng"]:
                        detail_list.column(col, width=80, anchor="center")
                    elif col == "Tình Trạng":
                        detail_list.column(col, width=120)
                    else:
                        detail_list.column(col, width=250)

                # Add data to table
                for detail in details:
                    detail_list.insert("", "end", values=(
                        detail[0],  # MaChiTietTra
                        detail[1],  # MaSach
                        detail[2],  # TenSach
                        detail[3],  # SoLuong
                        detail[4]  # TinhTrangSach
                    ))

                # Add scrollbar
                detail_frame = tk.Frame(details_window)
                detail_frame.pack(fill="both", expand=True, padx=20, pady=10)

                detail_y_scrollbar = ttk.Scrollbar(detail_frame, orient="vertical", command=detail_list.yview)
                detail_list.configure(yscrollcommand=detail_y_scrollbar.set)

                detail_list.pack(side="left", fill="both", expand=True)
                detail_y_scrollbar.pack(side="right", fill="y")

                # Add print button
                btn_print = tk.Button(details_window, text="In phiếu",
                                      command=lambda: self.print_return_slip(ma_phieu_tra),
                                      bg="#9b59b6", fg="white", font=("Arial", 10), width=10)
                btn_print.pack(side="left", padx=10, pady=15)

                # Close button
                btn_close = tk.Button(details_window, text="Đóng", command=details_window.destroy,
                                      bg="#3498db", fg="white", font=("Arial", 10), width=10)
                btn_close.pack(side="right", padx=10, pady=15)

        return_history_list.bind("<Double-1>", show_return_history_details)

        # Handle selection in active borrows table
        def on_return_selected(event):
            selected_items = return_list.selection()
            if selected_items:
                selected_item = selected_items[0]
                values = return_list.item(selected_item, "values")

                self.selected_borrow = values

                # Update form fields
                entries_return["Mã Độc Giả"].set(values[1])

                entries_return["Mã Sách"].configure(state="normal")
                entries_return["Mã Sách"].delete(0, tk.END)
                entries_return["Mã Sách"].insert(0, values[3])
                entries_return["Mã Sách"].configure(state="readonly")

                entries_return["Tên Sách"].configure(state="normal")
                entries_return["Tên Sách"].delete(0, tk.END)
                entries_return["Tên Sách"].insert(0, values[4])
                entries_return["Tên Sách"].configure(state="readonly")

                # Set default quantity as remaining
                entries_return["Số Lượng"].delete(0, tk.END)
                entries_return["Số Lượng"].insert(0, values[7])  # Remaining quantity

        return_list.bind("<<TreeviewSelect>>", on_return_selected)

        # Load initial data
        refresh_borrow_list()
        refresh_active_borrows()
        load_return_history()

    def show_users_management(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        lbl_title = tk.Label(self.content_frame, text="Quản lý người dùng", font=("Arial", 16, "bold"), bg="#f8f9fa")
        lbl_title.pack(pady=10)

        # Create frames
        form_frame = tk.Frame(self.content_frame, bg="#f8f9fa")
        form_frame.pack(pady=10, fill="x", padx=20)

        # Input fields
        fields = ["Mã Độc Giả", "Họ và Tên", "Giới Tính", "Ngày Sinh", "Điện Thoại", "Mật Khẩu", "Phân Quyền"]
        entries = {}

        # Number of fields
        num_fields = len(fields)
        mid_index = (num_fields + 1) // 2

        for i, field_text in enumerate(fields):
            # Create frame for each field
            field_frame = tk.Frame(form_frame, bg="#f8f9fa")
            field_frame.grid(row=i // mid_index, column=i % mid_index, padx=10, pady=10, sticky="w")

            # Add label
            label = tk.Label(field_frame, text=field_text, font=("Arial", 10), bg="#f8f9fa", width=12, anchor="w")
            label.pack(side="left")

            if field_text == "Giới Tính":
                entry = ttk.Combobox(field_frame, values=["Nam", "Nữ", "Khác"], width=25)
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Phân Quyền":
                entry = ttk.Combobox(field_frame, values=["Giảng viên", "Sinh viên"], width=25)
                entry.pack(side="left", fill="x", expand=True)
            elif field_text == "Mật Khẩu":
                entry = tk.Entry(field_frame, width=25, show="•")
                entry.pack(side="left", fill="x", expand=True)
            else:
                entry = tk.Entry(field_frame, width=25)
                entry.pack(side="left", fill="x", expand=True)
            entries[field_text] = entry
            # Variable to store selected user
        self.selected_user = None

        # Function to clear fields
        def clear_entries():
            for entry in entries.values():
                if hasattr(entry, 'delete'):
                    entry.delete(0, tk.END)
                elif hasattr(entry, 'set'):
                    entry.set('')
            self.selected_user = None

        # Function to refresh user list
        def refresh_user_list():
            for item in user_list.get_children():
                user_list.delete(item)

            users = self.user_manager.get_all_users()
            for user in users:
                user_list.insert("", "end", values=(
                    user[0],  # MaDocGia
                    user[1],  # HoVaTen
                    user[2],  # GioiTinh
                    user[3],  # NgaySinh
                    user[4],  # DienThoai
                    user[6]  # PhanQuyen
                ))

        # Button frame
        button_frame = tk.Frame(self.content_frame, bg="#f8f9fa")
        button_frame.pack(pady=10)

        # Add user function
        def add_user():
            try:
                ma_doc_gia = entries["Mã Độc Giả"].get().strip()
                ho_ten = entries["Họ và Tên"].get().strip()
                gioi_tinh = entries["Giới Tính"].get().strip()
                ngay_sinh = entries["Ngày Sinh"].get().strip()
                dien_thoai = entries["Điện Thoại"].get().strip()
                mat_khau = entries["Mật Khẩu"].get().strip()
                phan_quyen = entries["Phân Quyền"].get().strip()

                # Validate data
                if not all([ma_doc_gia, ho_ten, gioi_tinh, ngay_sinh, dien_thoai, mat_khau, phan_quyen]):
                    messagebox.showwarning("Cảnh báo", "Vui lòng điền đầy đủ thông tin!")
                    return

                # Validate date format
                try:
                    datetime.strptime(ngay_sinh, "%Y-%m-%d")
                except ValueError:
                    messagebox.showwarning("Cảnh báo", "Ngày sinh không hợp lệ! Định dạng: YYYY-MM-DD")
                    return

                # Add user to database
                if self.user_manager.add_user(ma_doc_gia, ho_ten, gioi_tinh, ngay_sinh, dien_thoai, mat_khau,
                                              phan_quyen):
                    messagebox.showinfo("Thành công", "Thêm người dùng thành công!")
                    clear_entries()
                    refresh_user_list()
                else:
                    messagebox.showerror("Lỗi", "Không thể thêm người dùng. Mã độc giả có thể đã tồn tại.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Edit user function
        def edit_user():
            if not self.selected_user:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn người dùng cần sửa!")
                return

            try:
                ma_doc_gia = entries["Mã Độc Giả"].get().strip()
                ho_ten = entries["Họ và Tên"].get().strip()
                gioi_tinh = entries["Giới Tính"].get().strip()
                ngay_sinh = entries["Ngày Sinh"].get().strip()
                dien_thoai = entries["Điện Thoại"].get().strip()
                mat_khau = entries["Mật Khẩu"].get().strip()
                phan_quyen = entries["Phân Quyền"].get().strip()  # Lấy phân quyền

                # Validate data
                if not all([ma_doc_gia, ho_ten, gioi_tinh, ngay_sinh, dien_thoai, phan_quyen]):
                    messagebox.showwarning("Cảnh báo", "Vui lòng điền đầy đủ thông tin!")
                    return

                # Validate date format
                try:
                    datetime.strptime(ngay_sinh, "%Y-%m-%d")
                except ValueError:
                    messagebox.showwarning("Cảnh báo", "Ngày sinh không hợp lệ! Định dạng: YYYY-MM-DD")
                    return

                # If password is empty, keep the old password
                if not mat_khau:
                    user = self.user_manager.get_user_by_id(ma_doc_gia)
                    if user:
                        mat_khau = user[5]  # Get current password

                # Update user
                if self.user_manager.update_user(ma_doc_gia, ho_ten, gioi_tinh, ngay_sinh, dien_thoai, mat_khau,
                                                 phan_quyen):
                    messagebox.showinfo("Thành công", "Cập nhật thông tin người dùng thành công!")
                    clear_entries()
                    refresh_user_list()
                else:
                    messagebox.showerror("Lỗi", "Không thể cập nhật thông tin người dùng.")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Delete user function
        def delete_user():
            if not self.selected_user:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn người dùng cần xóa!")
                return

            if messagebox.askyesno("Xác nhận",
                                   f"Bạn có chắc chắn muốn xóa người dùng '{self.selected_user[1]}' không?"):
                try:
                    if self.user_manager.delete_user(self.selected_user[0]):
                        messagebox.showinfo("Thành công", "Xóa người dùng thành công!")
                        clear_entries()
                        refresh_user_list()
                    else:
                        messagebox.showerror("Lỗi", "Không thể xóa người dùng. Người dùng có thể đang mượn sách.")
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

        # Search user function
        def search_user():
            search_term = entries["Họ và Tên"].get().strip()
            if not search_term:
                refresh_user_list()
                return

            for item in user_list.get_children():
                user_list.delete(item)

            users = self.user_manager.search_users(search_term)
            for user in users:
                user_list.insert("", "end", values=(
                    user[0],  # MaDocGia
                    user[1],  # HoVaTen
                    user[2],  # GioiTinh
                    user[3],  # NgaySinh
                    user[4],  # DienThoai
                    user[6]  # PhanQuyen
                ))

        # Function buttons
        btn_add = tk.Button(button_frame, text="Thêm", bg="#1abc9c", fg="white", font=("Arial", 10), width=10,
                            command=add_user)
        btn_add.pack(side="left", padx=5)

        btn_edit = tk.Button(button_frame, text="Sửa", bg="#f39c12", fg="white", font=("Arial", 10), width=10,
                             command=edit_user)
        btn_edit.pack(side="left", padx=5)

        btn_delete = tk.Button(button_frame, text="Xóa", bg="#e74c3c", fg="white", font=("Arial", 10), width=10,
                               command=delete_user)
        btn_delete.pack(side="left", padx=5)

        btn_search = tk.Button(button_frame, text="Tìm kiếm", bg="#3498db", fg="white", font=("Arial", 10), width=10,
                               command=search_user)
        btn_search.pack(side="left", padx=5)

        btn_clear = tk.Button(button_frame, text="Làm mới", bg="#95a5a6", fg="white", font=("Arial", 10), width=10,
                              command=clear_entries)
        btn_clear.pack(side="left", padx=5)

        # Create user list table
        columns = ["Mã Độc Giả", "Họ và Tên", "Giới Tính", "Ngày Sinh", "Điện Thoại", "Phân Quyền"]
        user_list = ttk.Treeview(self.content_frame, columns=columns, show="headings", height=15)

        for col in columns:
            user_list.heading(col, text=col)
            if col == "Mã Độc Giả":
                user_list.column(col, width=80, anchor="center")
            elif col in ["Giới Tính", "Phân Quyền"]:
                user_list.column(col, width=100, anchor="center")
            elif col == "Ngày Sinh":
                user_list.column(col, width=100, anchor="center")
            elif col == "Điện Thoại":
                user_list.column(col, width=120)
            else:
                user_list.column(col, width=200)

        # Add scrollbar
        table_frame = tk.Frame(self.content_frame)
        table_frame.pack(fill="both", expand=True, padx=20, pady=10)

        y_scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=user_list.yview)
        user_list.configure(yscrollcommand=y_scrollbar.set)

        user_list.pack(side="left", fill="both", expand=True)
        y_scrollbar.pack(side="right", fill="y")

        # Handle row selection
        def on_user_selected(event):
            selected_items = user_list.selection()
            if selected_items:
                selected_item = selected_items[0]
                values = user_list.item(selected_item, "values")

                self.selected_user = values

                # Get detailed information from database (including password)
                user = self.user_manager.get_user_by_id(values[0])
                if user:
                    # Fill form fields
                    entries["Mã Độc Giả"].delete(0, tk.END)
                    entries["Mã Độc Giả"].insert(0, user[0])

                    entries["Họ và Tên"].delete(0, tk.END)
                    entries["Họ và Tên"].insert(0, user[1])

                    entries["Giới Tính"].set(user[2])

                    entries["Ngày Sinh"].delete(0, tk.END)
                    entries["Ngày Sinh"].insert(0, user[3])

                    entries["Điện Thoại"].delete(0, tk.END)
                    entries["Điện Thoại"].insert(0, user[4])

                    entries["Mật Khẩu"].delete(0, tk.END)
                    # Don't show actual password for security

                    entries["Phân Quyền"].set(user[6])

        user_list.bind("<<TreeviewSelect>>", on_user_selected)

        # Load initial data
        refresh_user_list()

    def show_reports(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        lbl_title = tk.Label(self.content_frame, text="Báo cáo thống kê", font=("Arial", 16, "bold"), bg="#f8f9fa")
        lbl_title.pack(pady=10)

        # Create notebook with report tabs
        tab_control = ttk.Notebook(self.content_frame)

        # Overview tab
        overview_tab = ttk.Frame(tab_control)
        tab_control.add(overview_tab, text='Tổng quan')

        # Popular books tab
        popular_books_tab = ttk.Frame(tab_control)
        tab_control.add(popular_books_tab, text='Sách phổ biến')

        # Overdue tab
        overdue_tab = ttk.Frame(tab_control)
        tab_control.add(overdue_tab, text='Sách quá hạn')

        tab_control.pack(expand=1, fill='both', padx=10, pady=5)

        # === OVERVIEW TAB ===
        overview_frame = tk.Frame(overview_tab, bg="#f8f9fa")
        overview_frame.pack(padx=20, pady=20, fill="both", expand=True)

        # Get statistics data
        book_stats = self.report_manager.get_book_statistics()
        user_stats = self.report_manager.get_user_statistics()
        borrow_stats = self.report_manager.get_borrow_statistics()

        # Books statistics display
        book_frame = tk.LabelFrame(overview_frame, text="Thống kê sách", font=("Arial", 12, "bold"), bg="#f8f9fa")
        book_frame.pack(pady=10, fill="x")

        if book_stats:
            total_books = book_stats[0] or 0
            total_copies = book_stats[1] or 0
            total_value = book_stats[2] or 0

            tk.Label(book_frame, text=f"Tổng số đầu sách: {total_books}", font=("Arial", 11), bg="#f8f9fa").pack(
                anchor="w", padx=10, pady=5)
            tk.Label(book_frame, text=f"Tổng số bản sách: {total_copies}", font=("Arial", 11), bg="#f8f9fa").pack(
                anchor="w", padx=10, pady=5)
            tk.Label(book_frame, text=f"Tổng giá trị: {total_value:,.2f} VND", font=("Arial", 11), bg="#f8f9fa").pack(
                anchor="w", padx=10, pady=5)

        # User statistics display
        user_frame = tk.LabelFrame(overview_frame, text="Thống kê người dùng", font=("Arial", 12, "bold"), bg="#f8f9fa")
        user_frame.pack(pady=10, fill="x")

        total_users = 0
        for user_type in user_stats:
            phan_quyen = user_type[0]
            count = user_type[1]
            total_users += count
            tk.Label(user_frame, text=f"Số lượng {phan_quyen}: {count}", font=("Arial", 11), bg="#f8f9fa").pack(
                anchor="w", padx=10, pady=5)

        tk.Label(user_frame, text=f"Tổng số người dùng: {total_users}", font=("Arial", 11, "bold"), bg="#f8f9fa").pack(
            anchor="w", padx=10, pady=5)

        # Borrow statistics display
        borrow_frame = tk.LabelFrame(overview_frame, text="Thống kê mượn sách", font=("Arial", 12, "bold"),
                                     bg="#f8f9fa")
        borrow_frame.pack(pady=10, fill="x")

        if borrow_stats:
            total_borrows = borrow_stats[0] or 0
            total_borrowers = borrow_stats[1] or 0
            total_borrowed_books = borrow_stats[2] or 0

            tk.Label(borrow_frame, text=f"Tổng số phiếu mượn: {total_borrows}", font=("Arial", 11), bg="#f8f9fa").pack(
                anchor="w", padx=10, pady=5)
            tk.Label(borrow_frame, text=f"Số người mượn: {total_borrowers}", font=("Arial", 11), bg="#f8f9fa").pack(
                anchor="w", padx=10, pady=5)
            tk.Label(borrow_frame, text=f"Tổng số sách đã mượn: {total_borrowed_books}", font=("Arial", 11),
                     bg="#f8f9fa").pack(anchor="w", padx=10, pady=5)

        # === POPULAR BOOKS TAB ===
        popular_frame = tk.Frame(popular_books_tab, bg="#f8f9fa")
        popular_frame.pack(padx=20, pady=20, fill="both", expand=True)

        tk.Label(popular_frame, text="Top 10 sách được mượn nhiều nhất", font=("Arial", 12, "bold"), bg="#f8f9fa").pack(
            pady=10)

        # Create popular books table
        columns = ["Mã Sách", "Tên Sách", "Tác Giả", "Thể Loại", "Số lần mượn"]
        popular_list = ttk.Treeview(popular_frame, columns=columns, show="headings", height=15)

        for col in columns:
            popular_list.heading(col, text=col)
            if col == "Mã Sách":
                popular_list.column(col, width=80, anchor="center")
            elif col == "Số lần mượn":
                popular_list.column(col, width=100, anchor="center")
            elif col == "Thể Loại":
                popular_list.column(col, width=120)
            elif col == "Tác Giả":
                popular_list.column(col, width=150)
            else:
                popular_list.column(col, width=250)

        # Add scrollbar
        popular_frame_table = tk.Frame(popular_frame)
        popular_frame_table.pack(fill="both", expand=True)

        popular_y_scrollbar = ttk.Scrollbar(popular_frame_table, orient="vertical", command=popular_list.yview)
        popular_list.configure(yscrollcommand=popular_y_scrollbar.set)

        popular_list.pack(side="left", fill="both", expand=True)
        popular_y_scrollbar.pack(side="right", fill="y")

        # Get popular books data
        popular_books = self.report_manager.get_popular_books(10)
        for book in popular_books:
            popular_list.insert("", "end", values=(
                book[0],  # MaSach
                book[1],  # TenSach
                book[2],  # TacGia
                book[3],  # TheLoai
                book[4]  # TimesIssued
            ))

        # === OVERDUE TAB ===
        overdue_frame = tk.Frame(overdue_tab, bg="#f8f9fa")
        overdue_frame.pack(padx=20, pady=20, fill="both", expand=True)

        tk.Label(overdue_frame, text="Danh sách sách quá hạn trả", font=("Arial", 12, "bold"), bg="#f8f9fa").pack(
            pady=10)

        # Create overdue books table
        columns = ["Mã Phiếu", "Mã Độc Giả", "Tên Độc Giả", "Mã Sách", "Tên Sách", "Số Lượng", "Ngày Mượn",
                   "Ngày Hẹn Trả", "Số Ngày Quá Hạn"]
        overdue_list = ttk.Treeview(overdue_frame, columns=columns, show="headings", height=15)

        for col in columns:
            overdue_list.heading(col, text=col)
            if col in ["Mã Phiếu", "Mã Độc Giả", "Mã Sách", "Số Lượng", "Số Ngày Quá Hạn"]:
                overdue_list.column(col, width=80, anchor="center")
            elif col in ["Ngày Mượn", "Ngày Hẹn Trả"]:
                overdue_list.column(col, width=100, anchor="center")
            else:
                overdue_list.column(col, width=150)

        # Add scrollbar
        overdue_frame_table = tk.Frame(overdue_frame)
        overdue_frame_table.pack(fill="both", expand=True)

        overdue_x_scrollbar = ttk.Scrollbar(overdue_frame_table, orient="horizontal", command=overdue_list.xview)
        overdue_list.configure(xscrollcommand=overdue_x_scrollbar.set)

        overdue_y_scrollbar = ttk.Scrollbar(overdue_frame_table, orient="vertical", command=overdue_list.yview)
        overdue_list.configure(yscrollcommand=overdue_y_scrollbar.set)

        overdue_list.pack(side="top", fill="both", expand=True)
        overdue_y_scrollbar.pack(side="right", fill="y")
        overdue_x_scrollbar.pack(side="bottom", fill="x")

        # Get overdue books data
        overdue_books = self.report_manager.get_overdue_books()
        for book in overdue_books:
            overdue_list.insert("", "end", values=(
                book[0],  # MaPhieuMuon
                book[1],  # MaDocGia
                book[2],  # HoVaTen
                book[5],  # MaSach
                book[6],  # TenSach
                book[7],  # SoLuong
                book[3],  # NgayMuon
                book[4],  # NgayHenTra
                f"{int(book[9])} ngày"  # DaysOverdue
            ))

        btn_export_pdf = tk.Button(self.content_frame, text="Xuất báo cáo", bg="#3498db", fg="white",
                                   font=("Arial", 10), width=20,
                                   command=self.export_report_word)  # Sửa ở đây
        btn_export_pdf.pack(pady=10)

    def export_report_word(self):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        from tkinter import filedialog

        # Tạo hộp thoại để lưu file
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            title="Lưu báo cáo thống kê"
        )

        if not file_path:  # Người dùng đã hủy việc lưu file
            return

        # Tạo tài liệu Word mới
        doc = Document()

        # Thiết lập style cho tiêu đề
        title = doc.add_heading('BÁO CÁO THỐNG KÊ THƯ VIỆN', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Thêm thời gian xuất báo cáo
        current_time = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        time_paragraph = doc.add_paragraph(f"Thời gian xuất báo cáo: {current_time}")
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # Thêm dòng trống

        # Thêm thống kê sách
        doc.add_heading('Thống kê sách', level=1)
        book_stats = self.report_manager.get_book_statistics()
        if book_stats:
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'

            # Thiết lập tiêu đề hàng
            heading_cells = table.rows[0].cells
            heading_cells[0].text = "Thông tin"
            heading_cells[1].text = "Giá trị"

            # Thêm dữ liệu
            data_rows = [
                ["Tổng số đầu sách", str(book_stats[0])],
                ["Tổng số bản sách", str(book_stats[1])],
                ["Tổng giá trị", f"{book_stats[2]:,.2f} VND"]
            ]

            for i, row_data in enumerate(data_rows):
                row = table.rows[i + 1].cells
                row[0].text = row_data[0]
                row[1].text = row_data[1]

        doc.add_paragraph()  # Thêm dòng trống

        # Thêm thống kê người dùng
        doc.add_heading('Thống kê người dùng', level=1)
        user_stats = self.report_manager.get_user_statistics()

        table = doc.add_table(rows=len(user_stats) + 2, cols=2)  # +1 cho hàng tiêu đề, +1 cho tổng
        table.style = 'Table Grid'

        # Thiết lập tiêu đề hàng
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Loại người dùng"
        heading_cells[1].text = "Số lượng"

        # Thêm dữ liệu người dùng
        total_users = 0
        for i, user_type in enumerate(user_stats):
            row = table.rows[i + 1].cells
            row[0].text = user_type[0]
            row[1].text = str(user_type[1])
            total_users += user_type[1]

        # Thêm dòng tổng
        last_row = table.rows[len(user_stats) + 1].cells
        last_row[0].text = "Tổng số người dùng"
        last_row[1].text = str(total_users)

        doc.add_paragraph()  # Thêm dòng trống

        # Thêm thống kê mượn sách
        doc.add_heading('Thống kê mượn sách', level=1)
        borrow_stats = self.report_manager.get_borrow_statistics()
        if borrow_stats:
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'

            # Thiết lập tiêu đề hàng
            heading_cells = table.rows[0].cells
            heading_cells[0].text = "Thông tin"
            heading_cells[1].text = "Giá trị"

            # Thêm dữ liệu
            data_rows = [
                ["Tổng số phiếu mượn", str(borrow_stats[0])],
                ["Số người mượn", str(borrow_stats[1])],
                ["Tổng số sách đã mượn", str(borrow_stats[2])]
            ]

            for i, row_data in enumerate(data_rows):
                row = table.rows[i + 1].cells
                row[0].text = row_data[0]
                row[1].text = row_data[1]

        doc.add_paragraph()  # Thêm dòng trống

        # Thêm sách phổ biến
        doc.add_heading('Top 10 sách được mượn nhiều nhất', level=1)
        popular_books = self.report_manager.get_popular_books(10)

        table = doc.add_table(rows=len(popular_books) + 1, cols=5)
        table.style = 'Table Grid'

        # Thiết lập tiêu đề hàng
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Mã Sách"
        heading_cells[1].text = "Tên Sách"
        heading_cells[2].text = "Tác Giả"
        heading_cells[3].text = "Thể Loại"
        heading_cells[4].text = "Số lần mượn"

        # Thêm dữ liệu sách phổ biến
        for i, book in enumerate(popular_books):
            row = table.rows[i + 1].cells
            row[0].text = str(book[0])
            row[1].text = book[1]
            row[2].text = book[2]
            row[3].text = book[3]
            row[4].text = str(book[4])

        # Thêm sách quá hạn
        doc.add_heading('Danh sách sách quá hạn trả', level=1)
        overdue_books = self.report_manager.get_overdue_books()

        table = doc.add_table(rows=len(overdue_books) + 1, cols=6)
        table.style = 'Table Grid'

        # Thiết lập tiêu đề hàng
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Mã Phiếu"
        heading_cells[1].text = "Mã Độc Giả"
        heading_cells[2].text = "Tên Độc Giả"
        heading_cells[3].text = "Mã Sách"
        heading_cells[4].text = "Tên Sách"
        heading_cells[5].text = "Số Ngày Quá Hạn"

        # Thêm dữ liệu sách quá hạn
        for i, book in enumerate(overdue_books):
            row = table.rows[i + 1].cells
            row[0].text = str(book[0])
            row[1].text = str(book[1])
            row[2].text = book[2]
            row[3].text = str(book[5])
            row[4].text = book[6]
            row[5].text = f"{int(book[9])} ngày"

        # Lưu file
        doc.save(file_path)

        # Thông báo thành công
        messagebox.showinfo("Thông báo", f"Đã xuất báo cáo Word thành công!\nFile được lưu tại: {file_path}")

    def logout(self):
            if messagebox.askyesno("Đăng xuất", "Bạn có chắc chắn muốn đăng xuất?"):
                self.root.destroy()
            import os
            import sys
            os.execl(sys.executable, sys.executable, *sys.argv)

    def on_close(self):
        if messagebox.askyesno("Thoát", "Bạn có chắc chắn muốn thoát khỏi ứng dụng?"):
            self.db.close_connection()
            self.root.destroy()